import os
import subprocess
from flask import Flask, render_template, request, send_file
from docx import Document

app = Flask(__name__)

@app.route('/note_de_frais', methods=['GET', 'POST'])
def note_de_frais():
    if request.method == 'POST':
        # Récupérer les données du formulaire
        nom = request.form['nom']
        date = request.form['date']
        ndc = request.form['ndc']
        date1 = request.form['date1']
        desc1 = request.form['desc1']
        mont1 = request.form['mont1']
        date2 = request.form.get('date2', '')
        desc2 = request.form.get('desc2', '')
        mont2 = request.form.get('mont2', '')
        date3 = request.form.get('date3', '')
        desc3 = request.form.get('desc3', '')
        mont3 = request.form.get('mont3', '')
        dater = request.form['dater']
        datep = request.form['datep']

        # Charger le modèle Word
        template_path = 'templates_word/ndf.docx'
        doc = Document(template_path)

        # Remplacer les placeholders dans tous les paragraphes
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace("{{nom}}", nom)
            paragraph.text = paragraph.text.replace("{{date}}", date)
            paragraph.text = paragraph.text.replace("{{ndc}}", ndc)
            paragraph.text = paragraph.text.replace("{{dater}}", dater)
            paragraph.text = paragraph.text.replace("{{datep}}", datep)
            paragraph.text = paragraph.text.replace("{{date1}}", date1).replace("{{desc1}}", desc1).replace("{{mont1}}", mont1)
            paragraph.text = paragraph.text.replace("{{date2}}", date2).replace("{{desc2}}", desc2).replace("{{mont2}}", mont2)
            paragraph.text = paragraph.text.replace("{{date3}}", date3).replace("{{desc3}}", desc3).replace("{{mont3}}", mont3)

        # Remplacer les placeholders dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace("{{nom}}", nom)
                    cell.text = cell.text.replace("{{date}}", date)
                    cell.text = cell.text.replace("{{ndc}}", ndc)
                    cell.text = cell.text.replace("{{dater}}", dater)
                    cell.text = cell.text.replace("{{datep}}", datep)
                    cell.text = cell.text.replace("{{date1}}", date1).replace("{{desc1}}", desc1).replace("{{mont1}}", mont1)
                    cell.text = cell.text.replace("{{date2}}", date2).replace("{{desc2}}", desc2).replace("{{mont2}}", mont2)
                    cell.text = cell.text.replace("{{date3}}", date3).replace("{{desc3}}", desc3).replace("{{mont3}}", mont3)

        # Sauvegarder le document rempli
        output_docx_path = 'exports/note_de_frais_rempli.docx'
        doc.save(output_docx_path)

        # Chemin vers l'exécutable LibreOffice
        libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.com"

        # Convertir le document Word en PDF avec LibreOffice
        output_pdf_path = output_docx_path.replace('.docx', '.pdf')
        try:
            subprocess.run([
                libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', 'exports',
                output_docx_path
            ], check=True)
        except subprocess.CalledProcessError as e:
            return f"Erreur lors de la conversion du document en PDF : {e}"

        # Retourner le fichier PDF
        return send_file(output_pdf_path, as_attachment=True)

    return render_template('note_de_frais.html')

if __name__ == "__main__":
    # Créer le dossier d'exports s'il n'existe pas
    os.makedirs('exports', exist_ok=True)
    app.run(debug=True)
